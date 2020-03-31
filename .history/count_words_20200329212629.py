import docx

document = docx.Document('/Users/nordenbox/Dropbox/Words/englishpath/millionbaozi.docx')

print(type(document))

print('总共有 {} 个段落'.format(len(document.paragraphs)))
para = document.paragraphs
text = para.text  # 每一段的文字内容
print(len(text))